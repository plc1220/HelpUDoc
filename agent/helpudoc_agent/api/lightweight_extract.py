"""Lightweight deterministic extraction for curated Knowledge uploads.

This module intentionally uses only format-native libraries already needed by
the on-demand document tools. It does not invoke an LLM, OCR, a vector store,
or a background processing pipeline.
"""
from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Any


_TEXT_SUFFIXES = {".csv", ".html", ".htm", ".json", ".md", ".tsv", ".txt"}


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    divider = ["---"] * width

    def render(row: list[str]) -> str:
        return "| " + " | ".join(cell.replace("|", "\\|").replace("\n", " ") for cell in row) + " |"

    return "\n".join([render(header), render(divider), *(render(row) for row in normalized[1:])])


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    sections: list[str] = [f"# {path.name}"]
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(f"## Page {page_number}\n\n{text}")
    return "\n\n".join(sections).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document  # type: ignore

    document = Document(str(path))
    sections: list[str] = [f"# {path.name}"]
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        style = str(getattr(paragraph.style, "name", "") or "").lower()
        if style.startswith("heading"):
            try:
                level = max(2, min(6, int(style.split()[-1]) + 1))
            except (TypeError, ValueError):
                level = 2
            sections.append(f"{'#' * level} {text}")
        else:
            sections.append(text)
    for table_number, table in enumerate(document.tables, start=1):
        rows = [[str(cell.text or "").strip() for cell in row.cells] for row in table.rows]
        rendered = _markdown_table(rows)
        if rendered:
            sections.append(f"## Table {table_number}\n\n{rendered}")
    return "\n\n".join(sections).strip()


def _extract_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if path.suffix.lower() not in {".csv", ".tsv"}:
        return f"# {path.name}\n\n{text}".strip()
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    return f"# {path.name}\n\n{_markdown_table(rows)}".strip()


def extract_workspace_document(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        markdown = _extract_pdf(path)
    elif suffix == ".docx":
        markdown = _extract_docx(path)
    elif suffix in _TEXT_SUFFIXES:
        markdown = _extract_text(path)
    else:
        raise ValueError(
            f"Unsupported Knowledge source {suffix or '[none]'}. "
            "Supported: PDF, DOCX, Markdown, text, CSV, TSV, JSON, and HTML."
        )
    summary = next(
        (
            line.strip()
            for line in markdown.splitlines()
            if line.strip() and not line.lstrip().startswith(("#", "|", "---"))
        ),
        path.name,
    )
    return {
        "title": path.name,
        "summary": summary[:500],
        "markdown": markdown,
    }
