"""Complete, locator-preserving PDF, DOCX, and text extraction adapters."""
from __future__ import annotations

import csv
import hashlib
import io
import os
import re
import unicodedata
from pathlib import Path

from .models import ExtractionManifest, ExtractionWarning, ModelUsageRecord, SourceBlock
from .ocr import OcrAdapter, normalize_ocr_mode


EXTRACTOR_VERSION = "helpudoc-extractor/3"
_TEXT_SUFFIXES = {".csv", ".html", ".htm", ".json", ".md", ".tsv", ".txt"}


def _content_hash(text: str) -> str:
    normalized = unicodedata.normalize("NFC", text)
    return "sha256:" + hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _block(block_id: str, ordinal: int, text: str, block_type: str, **location: object) -> SourceBlock:
    return SourceBlock(
        id=block_id,
        ordinal=ordinal,
        text=unicodedata.normalize("NFC", text).strip(),
        blockType=block_type,  # type: ignore[arg-type]
        contentHash=_content_hash(text),
        **location,
    )


def pdf_pages_requiring_ocr(
    path: Path,
    *,
    mode: str | None = None,
    text_threshold: int | None = None,
) -> list[int]:
    """Return pages selected by the explicit OCR policy without mutating extraction state."""
    selected_mode = normalize_ocr_mode(mode)
    if selected_mode == "off":
        return []
    threshold = max(0, int(text_threshold or os.environ.get("KNOWLEDGE_OCR_TEXT_THRESHOLD", "40")))
    try:
        import fitz  # type: ignore

        document = fitz.open(str(path))
        try:
            if selected_mode == "always":
                return list(range(1, document.page_count + 1))
            return [
                page_index + 1
                for page_index in range(document.page_count)
                if len(re.sub(r"\s+", "", document.load_page(page_index).get_text("text") or "")) < threshold
            ]
        finally:
            document.close()
    except ImportError:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        if selected_mode == "always":
            return list(range(1, len(reader.pages) + 1))
        return [
            page_number
            for page_number, page in enumerate(reader.pages, start=1)
            if len(re.sub(r"\s+", "", page.extract_text() or "")) < threshold
        ]


def preflight_document(path: Path, *, ocr_text_threshold: int | None = None) -> dict[str, object]:
    """Cheap source-size/cost inputs without OCR or semantic parsing calls."""
    threshold = max(0, int(ocr_text_threshold or os.environ.get("KNOWLEDGE_OCR_TEXT_THRESHOLD", "40")))
    if path.suffix.lower() == ".pdf":
        try:
            import fitz  # type: ignore

            document = fitz.open(str(path))
            native_characters = 0
            ocr_pages: list[int] = []
            try:
                for page_index in range(document.page_count):
                    text = document.load_page(page_index).get_text("text") or ""
                    characters = len(re.sub(r"\s+", "", text))
                    native_characters += characters
                    if characters < threshold:
                        ocr_pages.append(page_index + 1)
                source_units = document.page_count
            finally:
                document.close()
        except ImportError:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            texts = [page.extract_text() or "" for page in reader.pages]
            counts = [len(re.sub(r"\s+", "", text)) for text in texts]
            native_characters = sum(counts)
            ocr_pages = [index + 1 for index, count in enumerate(counts) if count < threshold]
            source_units = len(reader.pages)
        return {
            "sourceType": "pdf",
            "bytes": path.stat().st_size,
            "sourceUnits": source_units,
            "nativeCharacters": native_characters,
            "ocrSourceUnits": len(ocr_pages),
            "ocrPages": ocr_pages,
            "ocrTextThreshold": threshold,
        }
    return {
        "sourceType": path.suffix.lower().lstrip(".") or "document",
        "bytes": path.stat().st_size,
        "sourceUnits": None,
        "nativeCharacters": None,
        "ocrSourceUnits": 0,
        "ocrPages": [],
        "ocrTextThreshold": threshold,
    }


def _pdf_blocks(
    path: Path,
    ocr_adapter: OcrAdapter | None = None,
    *,
    ocr_mode: str | None = None,
    ocr_text_threshold: int | None = None,
) -> tuple[list[SourceBlock], ExtractionManifest]:
    warnings: list[ExtractionWarning] = []
    blocks: list[SourceBlock] = []
    selected_mode = normalize_ocr_mode(ocr_mode or getattr(ocr_adapter, "mode", None))
    text_threshold = max(0, int(ocr_text_threshold or os.environ.get("KNOWLEDGE_OCR_TEXT_THRESHOLD", "40")))
    processed_pages = 0
    failed_pages = 0
    needs_ocr_pages = 0
    page_count = 0

    def page_media_artifact(page_number: int) -> str | None:
        for artifact in getattr(ocr_adapter, "media_artifacts", []) or []:
            if int(artifact.get("page") or 0) == page_number:
                return str(artifact.get("id") or "") or None
        return None

    def finalize_page(
        page_number: int,
        native: list[SourceBlock],
    ) -> None:
        nonlocal processed_pages, failed_pages, needs_ocr_pages
        native_chars = len(re.sub(r"\s+", "", "".join(block.text for block in native)))
        should_ocr = selected_mode == "always" or (selected_mode == "auto" and native_chars < text_threshold)
        if not should_ocr and native:
            start_ordinal = len(blocks)
            blocks.extend(block.model_copy(update={"ordinal": start_ordinal + index}) for index, block in enumerate(native))
            processed_pages += 1
            if selected_mode == "off" and native_chars < text_threshold:
                needs_ocr_pages += 1
                warnings.append(ExtractionWarning(
                    sourceUnit=f"page:{page_number}",
                    code="ocr_disabled_native_partial",
                    message=(
                        f"Page has only {native_chars} native-text characters; OCR is disabled and the "
                        "partial native text was retained."
                    ),
                ))
            return
        if not should_ocr:
            needs_ocr_pages += 1
            failed_pages += 1
            warnings.append(ExtractionWarning(
                sourceUnit=f"page:{page_number}",
                code="ocr_disabled",
                message="Page contains insufficient native text and OCR is disabled.",
            ))
            return
        needs_ocr_pages += 1
        outcome_getter = getattr(ocr_adapter, "page_outcome", None)
        outcome = outcome_getter(page_number) if callable(outcome_getter) else None
        ocr_blocks = ocr_adapter.recognize_pdf_page(path, page_number) if ocr_adapter else []
        media_artifact_id = page_media_artifact(page_number)
        emitted_ocr = 0
        for ocr_index, ocr in enumerate(ocr_blocks, start=1):
            if not ocr.text.strip():
                continue
            if ocr.blockType == "page_number":
                continue
            blocks.append(_block(
                f"pdf-p{page_number}-ocr{ocr_index}",
                len(blocks),
                ocr.text,
                "heading" if ocr.blockType == "heading" else ocr.blockType,
                page=page_number,
                bbox=ocr.bbox,
                headingLevel=ocr.headingLevel,
                extractionMethod="ocr",
                extractionConfidence=ocr.confidence,
                mediaArtifactId=media_artifact_id,
            ))
            emitted_ocr += 1
        if emitted_ocr or (outcome is not None and outcome.status in {"completed", "cached"} and outcome.blank):
            processed_pages += 1
            return
        if native:
            start_ordinal = len(blocks)
            blocks.extend(block.model_copy(update={"ordinal": start_ordinal + index}) for index, block in enumerate(native))
            processed_pages += 1
            warnings.append(ExtractionWarning(
                sourceUnit=f"page:{page_number}",
                code="ocr_fallback_native",
                message=(
                    f"Gemini Flash Lite OCR failed; retained {native_chars} characters of native text. "
                    + (str(getattr(outcome, "error", "") or "") or "No OCR result was returned.")
                ),
            ))
            return
        failed_pages += 1
        code = "ocr_disabled" if selected_mode == "off" else "ocr_failed" if outcome else "needs_ocr"
        reason = str(getattr(outcome, "error", "") or "") if outcome else ""
        warnings.append(ExtractionWarning(
            sourceUnit=f"page:{page_number}",
            code=code,
            message=(
                "Page contains insufficient native text and OCR is disabled."
                if selected_mode == "off"
                else f"Gemini Flash Lite OCR did not produce citable text. {reason}".strip()
            ),
        ))
    try:
        import fitz  # type: ignore

        document = fitz.open(str(path))
        page_count = document.page_count
        for page_index in range(page_count):
            page_number = page_index + 1
            page = document.load_page(page_index)
            page_blocks = sorted(page.get_text("blocks"), key=lambda item: (round(item[1], 1), item[0]))
            native: list[SourceBlock] = []
            for block_index, item in enumerate(page_blocks, start=1):
                text = str(item[4] or "").strip()
                if not text:
                    continue
                block_id = f"pdf-p{page_number}-b{block_index}"
                native.append(_block(
                    block_id,
                    len(native),
                    text,
                    "paragraph",
                    page=page_number,
                    bbox=(float(item[0]), float(item[1]), float(item[2]), float(item[3])),
                ))
            finalize_page(page_number, native)
        document.close()
    except ImportError:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            native: list[SourceBlock] = []
            if text:
                native.append(_block(
                    f"pdf-p{page_number}-b1",
                    0,
                    text,
                    "paragraph",
                    page=page_number,
                    extractionMethod="fallback",
                    extractionConfidence=0.85,
                ))
            finalize_page(page_number, native)
    usage = [
        ModelUsageRecord.model_validate(item.model_dump(mode="json") if hasattr(item, "model_dump") else item)
        for item in (getattr(ocr_adapter, "usage", []) or [])
    ]
    manifest = ExtractionManifest(
        extractorVersion=EXTRACTOR_VERSION,
        sourceType="pdf",
        discoveredSourceUnits=page_count,
        processedSourceUnits=processed_pages,
        failedSourceUnits=failed_pages,
        needsOcrSourceUnits=needs_ocr_pages,
        ocrMode=selected_mode,
        ocrProvider=str(getattr(ocr_adapter, "provider", "") or "") or None,
        ocrModel=str(getattr(ocr_adapter, "model", "") or "") or None,
        ocrTextThreshold=text_threshold,
        modelUsage=usage,
        mediaArtifacts=list(getattr(ocr_adapter, "media_artifacts", []) or []),
        warnings=warnings,
    )
    return blocks, manifest


def _docx_blocks(path: Path) -> tuple[list[SourceBlock], ExtractionManifest]:
    from docx import Document  # type: ignore

    document = Document(str(path))
    blocks: list[SourceBlock] = []
    heading_path: list[str] = []
    discovered = len(document.paragraphs) + len(document.tables)
    processed = 0
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = str(paragraph.text or "").strip()
        if not text:
            processed += 1
            continue
        style = str(getattr(paragraph.style, "name", "") or "")
        match = re.match(r"heading\s+(\d+)", style, re.IGNORECASE)
        if match:
            level = max(1, min(6, int(match.group(1))))
            heading_path = heading_path[: level - 1] + [text]
            blocks.append(_block(
                f"docx-p{paragraph_index}", len(blocks), text, "heading",
                paragraph=paragraph_index, headingLevel=level, headingPath=list(heading_path),
            ))
        else:
            blocks.append(_block(
                f"docx-p{paragraph_index}", len(blocks), text, "paragraph",
                paragraph=paragraph_index, headingPath=list(heading_path),
            ))
        processed += 1
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[str(cell.text or "").strip() for cell in row.cells] for row in table.rows]
        table_lines: list[str] = []
        if rows:
            width = max(len(row) for row in rows)
            normalized = [row + [""] * (width - len(row)) for row in rows]
            render = lambda row: "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"
            table_lines = [render(normalized[0]), render(["---"] * width)]
            table_lines.extend(render(row) for row in normalized[1:])
        blocks.append(_block(
            f"docx-t{table_index}", len(blocks), "\n".join(table_lines), "table",
            table=table_index, headingPath=list(heading_path),
        ))
        processed += 1
    return blocks, ExtractionManifest(
        extractorVersion=EXTRACTOR_VERSION,
        sourceType="docx",
        discoveredSourceUnits=discovered,
        processedSourceUnits=processed,
        failedSourceUnits=0,
    )


def _text_blocks(path: Path) -> tuple[list[SourceBlock], ExtractionManifest]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    suffix = path.suffix.lower()
    blocks: list[SourceBlock] = []
    if suffix in {".csv", ".tsv"}:
        rows = list(csv.reader(io.StringIO(raw), delimiter="\t" if suffix == ".tsv" else ","))
        if rows:
            width = max(len(row) for row in rows)
            normalized = [row + [""] * (width - len(row)) for row in rows]
            render = lambda row: "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"
            rendered = [render(normalized[0]), render(["---"] * width)]
            rendered.extend(render(row) for row in normalized[1:])
            blocks.append(_block("text-table1", 0, "\n".join(rendered), "table", table=1))
        discovered = len(rows)
    else:
        units = [unit.strip() for unit in re.split(r"\n\s*\n", raw) if unit.strip()]
        for unit_index, text in enumerate(units, start=1):
            heading = re.match(r"^(#{1,6})\s+(.+)$", text)
            if heading and "\n" not in text:
                blocks.append(_block(
                    f"text-b{unit_index}", len(blocks), heading.group(2), "heading",
                    paragraph=unit_index, headingLevel=len(heading.group(1)),
                ))
            else:
                blocks.append(_block(
                    f"text-b{unit_index}", len(blocks), text, "paragraph", paragraph=unit_index,
                ))
        discovered = len(units)
    return blocks, ExtractionManifest(
        extractorVersion=EXTRACTOR_VERSION,
        sourceType=suffix.lstrip(".") or "text",
        discoveredSourceUnits=discovered,
        processedSourceUnits=discovered,
        failedSourceUnits=0,
    )


def extract_blocks(
    path: Path,
    *,
    ocr_adapter: OcrAdapter | None = None,
    ocr_mode: str | None = None,
    ocr_text_threshold: int | None = None,
) -> tuple[list[SourceBlock], ExtractionManifest]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_blocks(
            path,
            ocr_adapter=ocr_adapter,
            ocr_mode=ocr_mode,
            ocr_text_threshold=ocr_text_threshold,
        )
    if suffix == ".docx":
        return _docx_blocks(path)
    if suffix in _TEXT_SUFFIXES:
        return _text_blocks(path)
    raise ValueError(
        f"Unsupported Knowledge source {suffix or '[none]'}. "
        "Supported: PDF, DOCX, Markdown, text, CSV, TSV, JSON, and HTML."
    )
