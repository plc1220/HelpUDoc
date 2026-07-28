from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from helpudoc_agent.api.lightweight_extract import extract_workspace_document
from helpudoc_agent.api.routes.documents import _resolve_workspace_document


def test_extract_docx_preserves_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Renewal Policy", level=1)
    document.add_paragraph("Customers must provide 30 days notice.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Notice"
    table.cell(1, 0).text = "Enterprise"
    table.cell(1, 1).text = "30 days"
    document.save(source)

    result = extract_workspace_document(source)

    assert result["title"] == "policy.docx"
    assert "## Renewal Policy" in result["markdown"]
    assert "Customers must provide 30 days notice." in result["markdown"]
    assert "| Tier | Notice |" in result["markdown"]


def test_extract_pdf_without_text_uses_filename_as_summary(tmp_path: Path) -> None:
    source = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with source.open("wb") as handle:
        writer.write(handle)

    result = extract_workspace_document(source)

    assert result == {
        "title": "blank.pdf",
        "summary": "blank.pdf",
        "markdown": "# blank.pdf",
    }


def test_extract_csv_renders_markdown_table(tmp_path: Path) -> None:
    source = tmp_path / "renewals.csv"
    source.write_text("account,risk\nAcme,high\n", encoding="utf-8")

    result = extract_workspace_document(source)

    assert "| account | risk |" in result["markdown"]
    assert "| Acme | high |" in result["markdown"]


def test_resolve_workspace_document_stays_inside_storage_root(tmp_path: Path) -> None:
    workspace_root = tmp_path / "workspaces"
    document = workspace_root / "workspace-1" / "notes.md"
    document.parent.mkdir(parents=True)
    document.write_text("# Notes", encoding="utf-8")

    assert _resolve_workspace_document(
        workspace_root,
        "workspace-1",
        "notes.md",
    ) == document.resolve()

    with pytest.raises(ValueError, match="workspaceId must remain inside"):
        _resolve_workspace_document(workspace_root, "../outside", "secret.txt")


def test_resolve_workspace_document_rejects_relative_path_traversal(tmp_path: Path) -> None:
    workspace_root = tmp_path / "workspaces"
    (workspace_root / "workspace-1").mkdir(parents=True)

    with pytest.raises(ValueError, match="Path must remain inside"):
        _resolve_workspace_document(workspace_root, "workspace-1", "../workspace-2/secret.txt")
