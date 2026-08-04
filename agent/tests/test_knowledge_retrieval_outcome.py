from __future__ import annotations

import json

from docx import Document

from helpudoc_agent.knowledge_ingestion.pipeline import extract_and_plan_document
from helpudoc_agent.state import WorkspaceState
from helpudoc_agent.tools.workspace.builtins.knowledge_navigation import (
    build_knowledge_navigation_tools,
)


def _concept(title: str, body: str, relationships: str = "", page: int = 1) -> str:
    return (
        "---\n"
        f'type: "Concept"\ntitle: "{title}"\n'
        "sources:\n"
        "  - id: \"source-span-1\"\n"
        "    locator:\n"
        "      kind: \"pdf_page_range\"\n"
        f"      start: {page}\n"
        f"      end: {page}\n"
        "---\n\n"
        f"# {title}\n\n{body}\n{relationships}"
)


def test_real_docx_processing_preserves_late_evidence_and_complete_core_ownership(tmp_path) -> None:
    source = tmp_path / "subscription-handbook.docx"
    document = Document()
    document.add_heading("Subscription Handbook", level=1)
    document.add_heading("Renewal Policy", level=2)
    for clause in range(1, 61):
        text = f"Policy clause {clause} applies to subscription renewals."
        if clause == 60:
            text += " Late evidence requires a 30-day renewal notice."
        document.add_paragraph(text)
    document.add_heading("Operational Controls", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = "Renewal notice"
    table.cell(1, 1).text = "Billing Service"
    document.save(source)

    plan = extract_and_plan_document(source)
    core_ids = [block_id for window in plan.windows for block_id in window.coreBlockIds]

    assert plan.manifest.failedSourceUnits == 0
    assert plan.manifest.processedSourceUnits == plan.manifest.discoveredSourceUnits
    assert len(core_ids) == len(plan.blocks)
    assert len(core_ids) == len(set(core_ids))
    assert set(core_ids) == {block.id for block in plan.blocks}
    assert "Late evidence requires a 30-day renewal notice" in plan.markdown
    assert any(node.title == "Renewal Policy" for node in plan.structure)
    assert any(block.blockType == "table" and "Billing Service" in block.text for block in plan.blocks)


def test_current_okf_snapshot_supports_lexical_graph_and_cited_read_retrieval(tmp_path) -> None:
    knowledge_root = tmp_path / ".system" / "knowledge" / "42"
    current = knowledge_root / "bundles" / "snapshot-good"
    policies = current / "concepts" / "policies"
    systems = current / "concepts" / "systems"
    risks = current / "concepts" / "risks"
    requirements = current / "concepts" / "requirements"
    policies.mkdir(parents=True)
    systems.mkdir(parents=True)
    risks.mkdir(parents=True)
    requirements.mkdir(parents=True)

    (current / "index.md").write_text("# Subscription Handbook\n", encoding="utf-8")
    (policies / "automatic-renewal-policy.md").write_text(
        _concept(
            "Automatic Renewal Policy",
            "Customers must receive a notice 30 days before renewal. "
            "The policy invokes the implementation engine.",
            "\n## Relationships\n\n"
            "* implemented by [Billing Service](../systems/billing-service.md)\n"
            "* mitigates [Renewal Risk](../risks/renewal-risk.md)\n",
            page=2,
        ),
        encoding="utf-8",
    )
    (systems / "billing-service.md").write_text(
        _concept(
            "Billing Service",
            "Schedules renewal notices and charges.",
            "\n## Relationships\n\n"
            "* dispatches through [Notification Queue](notification-queue.md)\n",
            page=3,
        ),
        encoding="utf-8",
    )
    (systems / "notification-queue.md").write_text(
        _concept("Notification Queue", "Delivers scheduled customer notices.", page=5),
        encoding="utf-8",
    )
    (risks / "renewal-risk.md").write_text(
        _concept("Renewal Risk", "Risk of renewal without adequate notice.", page=4),
        encoding="utf-8",
    )
    (requirements / "late-audit-requirement.md").write_text(
        _concept(
            "Late Audit Requirement",
            "The immutable audit ledger must retain renewal events for seven years.",
            page=73,
        ),
        encoding="utf-8",
    )

    stale = knowledge_root / "bundles" / "snapshot-old" / "concepts"
    stale.mkdir(parents=True)
    (stale.parent / "index.md").write_text("# Old bundle\n", encoding="utf-8")
    (stale / "obsolete.md").write_text(
        _concept("Obsolete Rule", "obsolete-secret-phrase should never be retrieved"),
        encoding="utf-8",
    )
    knowledge_root.mkdir(parents=True, exist_ok=True)
    (knowledge_root / "current.json").write_text(
        json.dumps(
            {
                "bundlePath": ".system/knowledge/42/bundles/snapshot-good",
                "snapshotHash": "snapshot-good",
            }
        ),
        encoding="utf-8",
    )

    workspace = WorkspaceState(workspace_id="retrieval-outcome", root_path=tmp_path)
    tools = {tool.name: tool for tool in build_knowledge_navigation_tools(workspace)}

    direct = json.loads(tools["knowledge_search"].invoke({"query": "30 days before renewal"}))
    assert direct["results"][0]["title"] == "Automatic Renewal Policy"
    assert "exact_phrase" in direct["results"][0]["reasons"]
    assert direct["results"][0]["snapshotId"] == "snapshot-good"

    late = json.loads(
        tools["knowledge_search"].invoke(
            {"query": "immutable audit ledger must retain renewal events"}
        )
    )
    assert late["results"][0]["title"] == "Late Audit Requirement"
    late_read = json.loads(
        tools["knowledge_read"].invoke(
            {"path": late["results"][0]["path"], "start_line": 1, "end_line": 40}
        )
    )
    assert late_read["sourceLocations"] == [{"pageStart": 73, "pageEnd": 73}]

    expanded = json.loads(tools["knowledge_search"].invoke({"query": "implementation engine"}))
    billing = next(result for result in expanded["results"] if result["title"] == "Billing Service")
    risk = next(result for result in expanded["results"] if result["title"] == "Renewal Risk")
    assert billing["reasons"] == ["graph"]
    assert risk["reasons"] == ["graph"]
    notification = next(
        result for result in expanded["results"] if result["title"] == "Notification Queue"
    )
    assert billing["graphHop"] == 1
    assert notification["reasons"] == ["graph"]
    assert notification["graphHop"] == 2
    assert billing["snapshotId"] == "snapshot-good"

    stale_search = json.loads(
        tools["knowledge_search"].invoke({"query": "obsolete-secret-phrase"})
    )
    assert stale_search["resultCount"] == 0

    absent = json.loads(
        tools["knowledge_search"].invoke({"query": "quantum banana payroll policy"})
    )
    assert absent["resultCount"] == 0

    read = json.loads(
        tools["knowledge_read"].invoke(
            {"path": billing["path"], "start_line": 1, "end_line": 40}
        )
    )
    assert read["title"] == "Billing Service"
    assert read["snapshotId"] == "snapshot-good"
    assert read["sourceLocations"] == [{"pageStart": 3, "pageEnd": 3}]
    assert "Schedules renewal notices" in read["content"]
